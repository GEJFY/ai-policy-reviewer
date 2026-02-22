"""DOCX document generator for revised text output."""

import io
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_revised_docx(
    title: str,
    original_text: str,
    revised_text: str,
    changes_applied: int,
    total_approved: int,
) -> io.BytesIO:
    """Generate a DOCX file containing the revised document text.

    Args:
        title: Document title
        original_text: Original document text
        revised_text: Text with approved suggestions applied
        changes_applied: Number of changes actually applied
        total_approved: Total number of approved findings

    Returns:
        BytesIO buffer containing the DOCX file
    """
    doc = DocxDocument()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Metadata
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"改訂版（{changes_applied}件の修正を適用）")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacing

    # Revised text - split by paragraphs
    for line in revised_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()
        run = para.add_run(stripped)
        run.font.size = Pt(11)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
